import os
import arrow
import wtforms
from io import BytesIO
from flask import render_template, request, redirect, url_for, flash, Response
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename

from docx import Document

from app import db
from . import project_bp as project
from .models import *
from .forms import *
from app.main.models import User
import requests
from pydrive.auth import ServiceAccountCredentials, GoogleAuth
from pydrive.drive import GoogleDrive

gauth = GoogleAuth()
keyfile_dict = requests.get(os.environ.get('GOOGLE_APPLICATION_CREDENTIALS')).json()
scopes = ['https://www.googleapis.com/auth/drive']
gauth.credentials = ServiceAccountCredentials.from_json_keyfile_dict(keyfile_dict, scopes)
drive = GoogleDrive(gauth)

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif'}
GANTT_ACTIVITIES = dict([(1, '1. พัฒนาโครงร่างการวิจัยและเครื่องมือการวิจัย'),
                         (2, '2. เสนอโครงร่างการวิจัยเพื่อขอรับการพิจารณาจริยธรรมฯ'),
                         (3, '3. เสนอขอรับทุนอุดหนุนการวิจัย'),
                         (4, '4. ผู้ทรงคุณวุฒิตรวจสอบและแก้ไข'),
                         (5, '5. ติดต่อประสานงานเพื่อขอเก็บข้อมูล'),
                         (6, '6. ดำเนินการเก็บรวบรวมข้อมูล'),
                         (7, '7. วิเคราะห์ผลการวิจัยและอภิปรายผล'),
                         (8, '8. จัดทำรายงานการวิจัยและเตรียมต้นฉบับตีพิมพ์งานวิจัย')]
                        )


def make_project_archive(project):
    """Create an archive of a project.
    :param project: a project model
    :return: archive model
    """

    archive = ProjectRecordArchive(
        project_record_id=project.id,
        archived_at=arrow.now(tz='Asia/Bangkok').datetime,
        title_th=project.title_th,
        subtitle_th=project.subtitle_th,
        title_en=project.title_en,
        subtitle_en=project.subtitle_en,
        objective=project.objective,
        method=project.method,
        intro=project.intro,
        status=project.status,
        prospected_journals=project.prospected_journals,
        use_applications=project.use_applications,
        created_at=project.created_at,
        updated_at=project.updated_at,
        creator_id=project.creator_id
    )
    members = []
    for member in project.members:
        if member.user:
            mem = {
                'role': member.role,
                'fullname_th': member.user.profile.fullname_th,
                'fullname_en': member.user.profile.fullname_en,
                'title_th': member.user.profile.title_th,
                'title_en': member.user.profile.title_en,
            }
        else:
            mem = {
                'role': member.role,
                'fullname_th': member.lastname,
                'fullname_en': member.firstname,
                'title_th': member.title,
                'title_en': '',
            }
        members.append(mem)
    figures = []
    for figure in project.figures:
        fig = {
            'title': figure.title,
            'desc': figure.desc,
            'fignum': figure.fignum,
            'url': figure.url,
        }
        figures.append(fig)
    milestones = []
    for milestone in project.milestones:
        mst = {
            'status': milestone.status,
            'detail': milestone.detail
        }
        milestones.append(mst)
    archive.members = members
    archive.figures = figures
    archive.milestones = milestones
    return archive


@project.route('/')
def list_projects():
    projects = ProjectRecord.query.filter_by(status='finished')
    return render_template('project/index.html', projects=projects)


@project.route('/user/<int:user_id>')
@login_required
def list_created_projects(user_id):
    user = User.query.get(user_id)
    projects = user.projects
    return render_template('project/created_projects.html', projects=projects)


@project.route('/detail/<int:project_id>')
@login_required
def display_project(project_id):
    project = ProjectRecord.query.get(project_id)

    if project.creator_id != current_user.id:
        return redirect(url_for('project.display_shorten', project_id=project.id))

    gantt_activities = []
    for a in sorted(project.gantt_activities, key=lambda x: x.task_id):
        gantt_activities.append([
            str(a.task_id),
            GANTT_ACTIVITIES.get(a.task_id),
            a.start_date.isoformat(),
            a.end_date.isoformat(),
            None, 0, None,
            a.start_date,
            a.end_date,
            a.id
        ])
    return render_template('project/detail.html', project=project, gantt_activities=gantt_activities)


@project.route('/detail/<int:project_id>/shorten')
@login_required
def display_project_shorten(project_id):
    project = ProjectRecord.query.get(project_id)
    return render_template('project/detail_shorten.html', project=project)


@project.route('/edit/<int:project_id>', methods=['GET', 'POST'])
@login_required
def edit_project(project_id):
    project = ProjectRecord.query.get(project_id)

    def edit_project_form_factory(project):
        class EditProjectRecordForm(ProjectRecordForm):
            parent = QuerySelectField(
                query_factory=lambda: ParentProjectRecord.query.all(),
                widget=wtforms.widgets.Select(),
                allow_blank=True,
                blank_text='โครงการเดี่ยว',
                default=project.parent_project
            )

        return EditProjectRecordForm

    EditProjectRecordForm = edit_project_form_factory(project)
    form = EditProjectRecordForm(obj=project)
    title = 'Edit Project Record'
    if request.method == 'POST':
        if form.validate_on_submit():
            form.populate_obj(project)
            project.updated_at = arrow.now(tz='Asia/Bangkok').datetime
            project.parent_project = form.parent.data

            if form.contract_upload.data:
                upfile = form.contract_upload.data
                filename = secure_filename(upfile.filename)
                upfile.save(filename)
                file_drive = drive.CreateFile({'title': filename})
                file_drive.SetContentFile(filename)
                file_drive.Upload()
                permission = file_drive.InsertPermission({'type': 'anyone',
                                                          'value': 'anyone',
                                                          'role': 'reader'})
                project.contract_url = file_drive['id']

            if form.final_report_upload.data:
                upfile = form.final_report_upload.data
                filename = secure_filename(upfile.filename)
                upfile.save(filename)
                file_drive = drive.CreateFile({'title': filename})
                file_drive.SetContentFile(filename)
                file_drive.Upload()
                permission = file_drive.InsertPermission({'type': 'anyone',
                                                          'value': 'anyone',
                                                          'role': 'reader'})
                project.final_report_url = file_drive['id']
            db.session.add(project)
            db.session.commit()
            flash('Data have been updated.', 'success')
        else:
            flash('Error occurred.', 'danger')
        return redirect(url_for('project.display_project', project_id=project.id))
    return render_template('project/project_edit.html', form=form, title=title, project=project)


@project.route('/parents/add', methods=['GET', 'POST'])
@login_required
def add_parent_project():
    form = ParentProjectRecordForm()
    if request.method == 'POST':
        if form.validate_on_submit():
            new_record = ParentProjectRecord()
            form.populate_obj(new_record)
            new_record.created_at = arrow.now(tz='Asia/Bangkok').datetime
            new_record.created_by = current_user
            db.session.add(new_record)
            db.session.commit()
            flash('A new parent project has been created.', 'success')
            return redirect(url_for('project.list_parent_projects'))
    return render_template('project/project_parent_add.html', form=form)


@project.route('/parents/<int:project_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_parent_project(project_id):
    record = ParentProjectRecord.query.get(project_id)
    form = ParentProjectRecordForm(obj=record)
    if request.method == 'POST':
        if form.validate_on_submit():
            form.populate_obj(record)
            db.session.add(record)
            db.session.commit()
            flash('A new parent project has been edited.', 'success')
            return redirect(url_for('project.list_parent_projects'))
        else:
            flash(form.errors, 'danger')
    return render_template('project/project_parent_add.html', form=form)


@project.route('/parents')
@login_required
def list_parent_projects():
    return render_template('project/parent_projects.html')


@project.route('/parents/<int:parent_id>/children')
@login_required
def list_children_projects(parent_id):
    parent = ParentProjectRecord.query.get(parent_id)
    return render_template('project/child_projects.html', parent=parent)


@project.route('/add', methods=['GET', 'POST'])
@login_required
def add_project():
    form = ProjectRecordForm()
    title = 'New Project Record'
    if request.method == 'POST':
        if form.validate_on_submit():
            new_proj = ProjectRecord()
            form.populate_obj(new_proj)
            new_proj.creator = current_user
            current_datetime = arrow.now(tz='Asia/Bangkok').datetime
            new_proj.created_at = current_datetime
            new_proj.updated_at = current_datetime
            new_proj.status = 'draft'
            db.session.add(new_proj)
            db.session.commit()
            flash('Data have been updated.', 'success')
        else:
            flash('Error occurred.', 'danger')
        return redirect(url_for('project.display_project', project_id=new_proj.id))
    return render_template('project/project_edit.html', form=form, title=title)


@project.route('/<int:project_id>/member/add', methods=['GET', 'POST'])
@login_required
def add_member(project_id):
    form = ProjectMemberForm()
    if form.validate_on_submit():
        new_member = ProjectMember(project_id=project_id)
        if form.users.data is None:
            form.populate_obj(new_member)
        else:
            new_member.user = form.users.data
            new_member.role = form.role.data
        db.session.add(new_member)
        db.session.commit()
        flash('New member has been added.', 'success')
        return redirect(url_for('project.display_project', project_id=project_id))
    return render_template('project/member_add.html', form=form)


@project.route('/<int:project_id>/member/<int:member_id>/remove', methods=['GET', 'POST'])
@login_required
def remove_member(project_id, member_id):
    member = ProjectMember.query.get(member_id)
    db.session.delete(member)
    db.session.commit()
    return redirect(url_for('project.display_project', project_id=project_id))


@project.route('/<int:project_id>/application/new', methods=['GET', 'POST'])
@login_required
def add_application(project_id):
    project = ProjectRecord.query.get(project_id)
    form = ApplicationForm()
    if request.method == 'POST':
        if form.validate_on_submit():
            new_app = Application()
            form.populate_obj(new_app)
            new_app.project = project
            db.session.add(new_app)
            db.session.commit()
            flash('Application has been recorded.', 'success')
        else:
            flash('Error occurred.', 'danger')
        return redirect(url_for('project.display_project', project_id=project.id))
    return render_template('project/application_add.html', form=form)


@project.route('/application/<int:app_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_application(app_id):
    app = Application.query.get(app_id)
    form = ApplicationForm(obj=app)
    if request.method == 'POST':
        if form.validate_on_submit():
            form.populate_obj(app)
            db.session.add(app)
            db.session.commit()
            flash('Application has been updated.', 'success')
        else:
            flash('Error occurred.', 'danger')
        return redirect(url_for('project.display_project', project_id=app.project.id))
    return render_template('project/application_add.html', form=form)


@project.route('/application/<int:app_id>/remove', methods=['GET', 'POST'])
@login_required
def remove_application(app_id):
    application = Application.query.get(app_id)
    project_id = application.project_id
    try:
        db.session.delete(application)
        db.session.commit()
    except:
        flash('Error occurred.', 'danger')
    else:
        flash('Application has been removed.', 'success')
    return redirect(url_for('project.display_project', project_id=project_id))


@project.route('/<int:project_id>/figure/add', methods=['GET', 'POST'])
@login_required
def add_figure(project_id):
    form = ProjectFigureForm()
    if request.method == 'POST':
        if form.validate_on_submit():
            if 'file' not in request.files:
                flash('No file uploaded.', 'warning')
                return redirect(request.url)
            file = request.files['file']
            if file.filename == '':
                flash('No selected file.', 'warning')
                return redirect(request.url)
            else:
                filename = secure_filename(file.filename)
                file.save(filename)
                file_drive = drive.CreateFile({'title': filename})
                file_drive.SetContentFile(filename)
                file_drive.Upload()
                permission = file_drive.InsertPermission({'type': 'anyone',
                                                          'value': 'anyone',
                                                          'role': 'reader'})
            new_figure = ProjectFigure()
            form.populate_obj(new_figure)
            new_figure.project_id = project_id
            new_figure.url = file_drive['id']
            db.session.add(new_figure)
            db.session.commit()
            return redirect(url_for('project.display_project', project_id=project_id))
    project = ProjectRecord.query.get(project_id)
    return render_template('project/figure_add.html', project=project, form=form)


@project.route('/figure/<int:figure_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_figure(figure_id):
    figure = ProjectFigure.query.get(figure_id)
    form = ProjectFigureForm(obj=figure)
    if request.method == 'POST':
        if form.validate_on_submit():
            form.populate_obj(figure)
            db.session.add(figure)
            db.session.commit()
            return redirect(url_for('project.display_project', project_id=figure.project.id))
    return render_template('project/figure_add.html', form=form)


@project.route('/figure/<int:figure_id>/remove', methods=['GET', 'POST'])
@login_required
def remove_figure(figure_id):
    figure = ProjectFigure.query.get(figure_id)
    try:
        file1 = drive.CreateFile({'id': figure.url})
        file1.Trash()
        db.session.delete(figure)
        db.session.commit()
    except:
        flash('ไม่สามารถลบรูปภาพได้', 'danger')
    else:
        flash('ลบไฟล์รูปภาพออกจากระบบเรียบร้อยแล้ว', 'success')
    return redirect(request.referrer)


@project.route('/<int:project_id>/submit', methods=['GET'])
@login_required
def submit_project(project_id):
    project = ProjectRecord.query.get(project_id)
    if len(project.members) == 0:
        flash('กรุณาเพิ่มรายชื่อนักวิจัยในโครงการก่อนยื่นขอพิจารณาโครงการ', 'warning')
        return redirect(url_for('project.display_project', project_id=project.id))

    if project.status == 'draft' or \
            project.status == 'concept revising':
        project.status = 'concept submitted'
    elif project.status == 'concept approved' or \
            project.status == 'full revising':
        project.status = 'full submitted'

    project.updated_at = arrow.now(tz='Asia/Bangkok').datetime
    project.submitted_at = arrow.now(tz='Asia/Bangkok').datetime
    db.session.add(project)
    archive = make_project_archive(project)
    db.session.add(archive)
    db.session.commit()
    flash('ยื่นขอพิจารณาโครงการเรียบร้อยแล้ว', 'success')
    return redirect(url_for('project.display_project', project_id=project.id))


@project.route('/<int:project_id>/ethic/add/confirm', methods=['GET', 'POST'])
@login_required
def confirm_add_ethic(project_id):
    project = ProjectRecord.query.get(project_id)
    return render_template('project/ethic_add.html', project=project)


@project.route('/<int:project_id>/ethic/add/confirmed', methods=['GET'])
@login_required
def add_ethic_request(project_id):
    project = ProjectRecord.query.get(project_id)
    ethic = ProjectEthicRecord(project_id=project_id,
                               submitted_at=arrow.now(tz='Asia/Bangkok').datetime,
                               status='submitted')
    archive = make_project_archive(project)
    project.updated_at = arrow.now(tz='Asia/Bangkok').datetime
    db.session.add(ethic)
    db.session.add(archive)
    db.session.add(project)
    db.session.commit()
    return redirect(url_for('project.display_project', project_id=project.id))


@project.route('/ethic/<int:ethic_id>/upload_file', methods=['GET', 'POST'])
@login_required
def upload_ethic_doc(ethic_id):
    ethic = ProjectEthicRecord.query.get(ethic_id)
    form = EthicRecordForm()
    if request.method == 'POST':
        if form.file_upload.data:
            upfile = form.file_upload.data
            filename = secure_filename(upfile.filename)
            upfile.save(filename)
            file_drive = drive.CreateFile({'title': filename})
            file_drive.SetContentFile(filename)
            file_drive.Upload()
            permission = file_drive.InsertPermission({'type': 'anyone',
                                                      'value': 'anyone',
                                                      'role': 'reader'})
            ethic.upload_file_url = file_drive['id']
            db.session.add(ethic)
            db.session.commit()
            flash('บันทึกไฟล์ที่อัพโหลดเรียบร้อยแล้ว', 'success')
            return redirect(url_for('project.display_project', project_id=ethic.project.id))
    return render_template('project/ethic_file_upload.html', form=form, ethic=ethic)


@project.route('/archive/<int:archive_id>', methods=['GET'])
@login_required
def view_archive(archive_id):
    ar = ProjectRecordArchive.query.get(archive_id)
    return render_template('project/archieve_detail.html', project=ar)


@project.route('/<int:project_id>/milestones/add', methods=['GET', 'POST'])
@login_required
def add_milestone(project_id):
    project = ProjectRecord.query.get(project_id)
    form = ProjectMilestoneForm()
    if request.method == 'POST':
        if form.validate_on_submit():
            new_milestone = ProjectMilestone()
            form.populate_obj(new_milestone)
            new_milestone.project_id = project_id
            new_milestone.created_at = arrow.now(tz='Asia/Bangkok').datetime
            db.session.add(new_milestone)
            db.session.commit()
            flash('New milestone added.', 'success')
            return redirect(url_for('project.display_project', project_id=project_id))
    return render_template('project/milestone_add.html', project=project, form=form)


@project.route('/<int:project_id>/milestones/<int:milestone_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_milestone(project_id, milestone_id):
    milestone = ProjectMilestone.query.get(milestone_id)
    project = ProjectRecord.query.get(project_id)
    form = ProjectMilestoneForm(obj=milestone)
    if request.method == 'POST':
        if form.validate_on_submit():
            form.populate_obj(milestone)
            db.session.add(milestone)
            db.session.commit()
            flash('The milestone has been update.', 'success')
            return redirect(url_for('project.display_project', project_id=project_id))
    return render_template('project/milestone_add.html', project=project, form=form)


@project.route('/<int:project_id>/milestones/clone')
@login_required
def clone_milestone(project_id):
    milestone = ProjectMilestone.query.filter_by(project_id=project_id) \
        .order_by(ProjectMilestone.created_at.desc()).first()
    if milestone:
        new_milestone = ProjectMilestone()
        new_milestone.project_id = milestone.project_id
        new_milestone.created_at = arrow.now(tz='Asia/Bangkok').datetime
        new_milestone.detail = milestone.detail
        new_milestone.status = milestone.status
        if milestone.gantt_activities:
            for a in milestone.gantt_activities:
                new_activity = ProjectGanttActivity()
                new_activity.created_at = arrow.now(tz='Asia/Bangkok').datetime
                new_activity.completion = a.completion
                new_activity.start_date = a.start_date
                new_activity.end_date = a.end_date
                new_activity.detail = a.detail
                new_activity.task_id = a.task_id
                new_milestone.gantt_activities.append(new_activity)
                db.session.add(new_activity)
        if milestone.budget_items:
            for b in milestone.budget_items:
                new_item = ProjectBudgetItem()
                new_item.sub_category = b.sub_category
                new_item.wage = b.wage
                new_item.phase = b.phase
                new_item.amount_spent = b.amount_spent
                new_item.created_at = arrow.now(tz='Asia/Bangkok').datetime
                new_item.edited_at = new_item.created_at
                new_milestone.budget_items.append(new_item)
        if milestone.summaries:
            for s in milestone.summaries:
                new_item = ProjectSummary()
                new_item.activity = s.activity
                new_item.expected_outcome = s.expected_outcome
                new_item.outcome = s.outcome
                new_item.created_at = arrow.now(tz='Asia/Bangkok').datetime
                new_item.edited_at = new_item.created_at
                new_milestone.summaries.append(new_item)
        db.session.add(new_milestone)
        db.session.commit()
        flash('New milestone added.', 'success')
    else:
        flash('No milestone to clone.', 'warning')
    return redirect(url_for('project.display_project', project_id=project_id))


@project.route('/<int:project_id>/milestones/<int:milestone_id>/gantt-activities')
@login_required
def list_gantt_activity(project_id, milestone_id):
    milestone = ProjectMilestone.query.get(milestone_id)
    gantt_activities = []
    for a in sorted(milestone.gantt_activities, key=lambda x: x.task_id):
        gantt_activities.append([
            str(a.task_id),
            GANTT_ACTIVITIES.get(a.task_id),
            a.start_date.isoformat(),
            a.end_date.isoformat(),
            None, float(a.completion), None
        ])
    return render_template('project/gantt_chart.html',
                           project_id=project_id,
                           milestone=milestone, gantt_activities=gantt_activities)


@project.route('/<int:project_id>/milestones/<int:milestone_id>/gantt-activities/add', methods=['GET', 'POST'])
@login_required
def add_gantt_activity(project_id, milestone_id):
    form = ProjectGanttActivityForm()
    if request.method == 'POST':
        if form.validate_on_submit():
            new_activity = ProjectGanttActivity()
            form.populate_obj(new_activity)
            new_activity.milestone_id = milestone_id
            new_activity.dirty = True
            new_activity.created_at = arrow.now(tz='Asia/Bangkok').datetime
            if new_activity.completion > 100 or new_activity.completion < 0:
                new_activity.completion = 100.0
            db.session.add(new_activity)
            db.session.commit()
            flash('New activity has been added.', 'success')
            return redirect(url_for('project.list_gantt_activity',
                                    project_id=project_id, milestone_id=milestone_id))
        else:
            flash(form.errors, 'danger')
    return render_template('project/gantt_activity_add.html', form=form,
                           project_id=project_id, milestone_id=milestone_id)


@project.route('projects/<int:project_id>/milestones/<int:milestone_id>/gantt-activities/<int:record_id>/edit',
               methods=['GET', 'POST'])
@login_required
def edit_gantt_activity(project_id, milestone_id, record_id):
    record = ProjectGanttActivity.query.get(record_id)
    form = ProjectGanttActivityForm(obj=record)
    if request.method == 'POST':
        if form.validate_on_submit():
            form.populate_obj(record)
            record.dirty = True
            db.session.add(record)
            db.session.commit()
            flash('Activity has been updated.', 'success')
            return redirect(url_for('project.list_gantt_activity',
                                    project_id=project_id,
                                    milestone_id=milestone_id))
        else:
            flash(form.errors, 'danger')
    return render_template('project/gantt_activity_edit.html',
                           project_id=project_id,
                           milestone=record.milestone,
                           form=form, record=record)


@project.route('projects/<int:project_id>/milestones/<int:milestone_id>/gantt-activities/<int:record_id>/remove')
@login_required
def remove_gantt_activity(project_id, milestone_id, record_id):
    record = ProjectGanttActivity.query.get(record_id)
    confirm = request.args.get('confirm', 'no')
    if record and confirm == 'yes':
        db.session.delete(record)
        db.session.commit()
        flash('The activity has been successfully removed.', 'success')
        return redirect(url_for('project.list_gantt_activity',
                                project_id=project_id, milestone_id=milestone_id))

    return render_template('project/gantt_remove.html', record=record,
                           project_id=project_id, milestone_id=milestone_id)


@project.route('projects/<int:project_id>/milestones/<int:milestone_id>/budgets')
@login_required
def list_budget_items(project_id, milestone_id):
    milestone = ProjectMilestone.query.get(milestone_id)
    return render_template('project/budgets.html',
                           milestone=milestone, project_id=project_id)


@project.route('projects/<int:project_id>/milestones/<int:milestone_id>/budgets/add', methods=['GET', 'POST'])
@login_required
def add_budget_item(project_id, milestone_id):
    form = ProjectBudgetItemForm()
    if request.method == 'POST':
        if form.validate_on_submit():
            new_item = ProjectBudgetItem()
            form.populate_obj(new_item)
            new_item.milestone_id = milestone_id
            new_item.created_at = arrow.now(tz='Asia/Bangkok').datetime
            new_item.edited_at = new_item.created_at
            db.session.add(new_item)
            db.session.commit()
            flash('New budget item has been added.', 'success')
            return redirect(url_for('project.list_budget_items',
                                    project_id=project_id, milestone_id=milestone_id))
        else:
            flash(form.errors, 'danger')
    return render_template('project/budget_item_add.html', form=form,
                           project_id=project_id, milestone_id=milestone_id)


@project.route('projects/<int:project_id>/milestones/<int:milestone_id>/budgets/<int:item_id>/edit',
               methods=['GET', 'POST'])
@login_required
def edit_budget_item(project_id, milestone_id, item_id):
    item = ProjectBudgetItem.query.get(item_id)
    form = ProjectBudgetItemForm(obj=item)
    if request.method == 'POST':
        if form.validate_on_submit():
            form.populate_obj(item)
            item.edited_at = arrow.now(tz='Asia/Bangkok').datetime
            db.session.add(item)
            db.session.commit()
            flash('The budget item has been updated.', 'success')
            return redirect(url_for('project.list_budget_items',
                                    project_id=project_id, milestone_id=milestone_id))
        else:
            flash(form.errors, 'danger')
    return render_template('project/budget_item_add.html', form=form, project_id=project_id, milestone_id=milestone_id)


@project.route('projects/<int:project_id>/milestones/<int:milestone_id>/summaries')
@login_required
def list_summaries(project_id, milestone_id):
    milestone = ProjectMilestone.query.get(milestone_id)
    return render_template('project/summary_records.html',
                           milestone=milestone, project_id=project_id)


@project.route('projects/<int:project_id>/milestones/<int:milestone_id>/summaries/add',
               methods=['GET', 'POST'])
@login_required
def add_summary_record(project_id, milestone_id):
    form = ProjectSummaryForm()
    if request.method == 'POST':
        if form.validate_on_submit():
            new_record = ProjectSummary()
            form.populate_obj(new_record)
            new_record.milestone_id = milestone_id
            new_record.edited_at = arrow.now(tz='Asia/Bangkok').datetime
            new_record.created_at = new_record.edited_at
            db.session.add(new_record)
            db.session.commit()
            flash('New summary record has been added.', 'success')
            return redirect(url_for('project.list_summaries',
                                    milestone_id=milestone_id,
                                    project_id=project_id))
        else:
            flash(form.errors, 'danger')
    return render_template('project/summary_record_edit.html', form=form,
                           project_id=project_id, milestone_id=milestone_id)


@project.route('projects/<int:project_id>/milestones/<int:milestone_id>/summaries/<int:record_id>/edit',
               methods=['GET', 'POST'])
@login_required
def edit_summary_record(project_id, milestone_id, record_id):
    record = ProjectSummary.query.get(record_id)
    form = ProjectSummaryForm(obj=record)
    if request.method == 'POST':
        if form.validate_on_submit():
            form.populate_obj(record)
            record.edited_at = arrow.now(tz='Asia/Bangkok').datetime
            db.session.add(record)
            db.session.commit()
            flash('New summary record has been added.', 'success')
            return redirect(url_for('project.list_summaries',
                                    milestone_id=milestone_id,
                                    project_id=project_id))
        else:
            flash(form.errors, 'danger')
    return render_template('project/summary_record_edit.html', form=form,
                           project_id=project_id, milestone_id=milestone_id)


@project.route('/admin/ethics', methods=['GET', 'POST'])
@login_required
def list_ethics(project_id):
    ethics = ProjectEthicRecord.query.get(project_id)
    return 'Hello'


@project.route('/journals/add', methods=['GET', 'POST'])
@login_required
def add_journal():
    form = ProjectJournalForm()
    project_id = request.args.get('project_id')
    if request.method == 'POST':
        if form.validate_on_submit():
            new_journal = ProjectPublicationJournal()
            form.populate_obj(new_journal)
            db.session.add(new_journal)
            db.session.commit()
            flash('New journal added.', 'success')
        else:
            flash(form.errors, 'danger')
        if project_id:
            return redirect(url_for('project.add_pub', project_id=project_id))
    return render_template('project/journal_add.html', form=form, project_id=project_id)


@project.route('/projects/<int:project_id>/pubs/<int:pub_id>/authors/add', methods=['GET', 'POST'])
@login_required
def add_pub_author(project_id, pub_id):
    form = ProjectPublicationAuthorForm()
    pub = ProjectPublication.query.get(pub_id)
    if request.method == 'POST':
        if form.validate_on_submit():
            new_author = ProjectPublicationAuthor()
            form.populate_obj(new_author)
            new_author.pub = pub
            new_author.user = form.users.data
            db.session.add(new_author)
            db.session.commit()
            flash('New author has been added.', 'success')
            return redirect(url_for('project.edit_pub_authors', pub_id=pub_id, project_id=project_id))
    return render_template('webadmin/pub_author_add.html', form=form)


@project.route('/<int:project_id>/pubs/add', methods=['GET', 'POST'])
@login_required
def add_pub(project_id):
    form = ProjectPublicationForm()
    if request.method == 'POST':
        if form.validate_on_submit():
            new_pub = ProjectPublication()
            form.populate_obj(new_pub)
            new_pub.project_id = project_id
            new_pub.journal = form.journals.data
            db.session.add(new_pub)
            db.session.commit()
            flash('New publication added.')
            return redirect(url_for('project.edit_pub_authors',
                                    project_id=project_id, pub_id=new_pub.id))
    return render_template('webadmin/pub_add.html', form=form, project_id=project_id)


@project.route('projects/<int:project_id>/pubs/<int:pub_id>/remove', methods=['GET', 'POST'])
@login_required
def remove_pub(project_id, pub_id):
    pub = ProjectPublication.query.get(pub_id)
    confirm = request.args.get('confirm', 'no')
    if confirm == 'yes':
        if pub:
            db.session.delete(pub)
            db.session.commit()
            flash('Data have been removed.', 'success')
            return redirect(url_for('project.display_project', project_id=project_id))
        else:
            flash('Publication no longer exists.', 'warning')
    return render_template('project/pub_remove.html', pub=pub, project_id=project_id)


@project.route('projects/<int:project_id>/pubs/<int:pub_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_pub(project_id, pub_id):
    pub = ProjectPublication.query.get(pub_id)
    form = ProjectPublicationForm(obj=pub)
    if request.method == 'POST':
        if form.validate_on_submit():
            form.populate_obj(pub)
            db.session.add(pub)
            db.session.commit()
            flash('Data have been saved.', 'success')
            return redirect(url_for('project.display_project', project_id=project_id))
        else:
            flash(form.errors, 'danger')
    return render_template('webadmin/pub_edit.html', form=form, pub=pub, project_id=project_id)


@project.route('projects/<int:project_id>/pubs/<int:pub_id>/authors/edit', methods=['GET', 'POST'])
@login_required
def edit_pub_authors(project_id, pub_id):
    pub = ProjectPublication.query.get(pub_id)
    return render_template('webadmin/pub_author_edit.html', pub=pub, project_id=project_id)


@project.route('/projects/<int:project_id>/pubs/<int:pub_id>/authors/<int:author_id>/remove')
@login_required
def remove_pub_author(project_id, pub_id, author_id):
    author = ProjectPublicationAuthor.query.get(author_id)
    if author:
        db.session.delete(author)
        db.session.commit()
        flash('Author has been deleted from publication', 'success')
    else:
        flash('The author with that ID was not found.', 'danger', )
    return redirect(url_for('project.edit_pub_authors', project_id=project_id, pub_id=pub_id))


@project.route('/<int:project_id>/pubs/<int:pub_id>/support/language/add', methods=['GET', 'POST'])
@login_required
def add_lang_support(project_id, pub_id):
    form = ProjectLanguageEditSupportForm()
    if request.method == 'POST':
        if form.validate_on_submit():
            new_support = ProjectLanguageEditingSupport()
            form.populate_obj(new_support)
            new_support.qualification = '|'.join(form.qualification_select.data)
            new_support.criteria = '|'.join(form.criteria_select.data)
            new_support.docs = '|'.join(form.docs_select.data)
            new_support.request = '|'.join(form.request_select.data)
            new_support.pub_id = pub_id
            new_support.submitted_at = arrow.now(tz='Asia/Bangkok').datetime,
            db.session.add(new_support)
            db.session.commit()
            flash('New language edit request added.')
            return redirect(url_for('project.display_project', project_id=project_id))
    return render_template('project/lang_support_add.html', project=project, form=form)


@project.route('/<int:project_id>/pubs/<int:pub_id>/support/language/<int:record_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_lang_support(project_id, pub_id, record_id):
    record = ProjectLanguageEditingSupport.query.get(record_id)
    form = ProjectLanguageEditSupportForm(obj=record)
    criteria = record.criteria.split('|')
    qualification = record.qualification.split('|')
    docs = record.docs.split('|')
    request_data = record.request.split('|')
    if request.method == 'POST':
        if form.validate_on_submit():
            form.populate_obj(record)
            record.qualification = '|'.join(form.qualification_select.data)
            record.criteria = '|'.join(form.criteria_select.data)
            record.docs = '|'.join(form.docs_select.data)
            record.request = '|'.join(form.request_select.data)
            record.pub_id = pub_id
            record.edited_at = arrow.now(tz='Asia/Bangkok').datetime,
            db.session.add(record)
            db.session.commit()
            flash('New language edit request edited.', 'success')
            return redirect(url_for('project.display_project', project_id=project_id))
        else:
            flash(form.errors, 'danger')
    return render_template('project/lang_support_edit.html',
                           project=project, form=form,
                           docs=docs, request_data=request_data,
                           qualification=qualification, criteria=criteria)


@project.route('/<int:project_id>/pubs/<int:pub_id>/support/reward/add', methods=['GET', 'POST'])
@login_required
def add_pub_reward(project_id, pub_id):
    pub = ProjectPublication.query.get(pub_id)
    form = ProjectPublishedRewardForm()
    if request.method == 'POST':
        if form.validate_on_submit():
            new_record = ProjectPublishedReward()
            form.populate_obj(new_record)
            new_record.pub_id = pub_id
            new_record.qualification = '|'.join(form.qualification_select.data)
            new_record.submitted_at = arrow.now(tz='Asia/Bangkok').datetime,
            if new_record.reward:
                reward_amt = float(new_record.reward.split()[-2])
            else:
                reward_amt = 0.0
            if new_record.apc:
                apc_amt = float(new_record.apc.split()[-2])
            else:
                apc_amt = 0.0
            new_record.amount = reward_amt + apc_amt
            db.session.add(new_record)
            db.session.commit()
            flash('Reward/Fee request has been added.', 'success')
            return redirect(url_for('project.display_project', project_id=project_id))
        else:
            flash(form.errors, 'danger')
    return render_template('project/reward_add.html',
                           pub=pub, project_id=project_id, form=form)


@project.route('/<int:project_id>/ethics/<int:ethic_id>')
@login_required
def view_ethic_record(project_id, ethic_id):
    ethic = ProjectEthicRecord.query.get(ethic_id)
    return render_template('project/ethic_review.html', ethic=ethic, project_id=project_id)


@project.route('/<int:project_id>/supports/proposal/add', methods=['GET', 'POST'])
@login_required
def add_proposal_development_support(project_id):
    form = ProjectProposalDevelopmentSupportForm()
    if request.method == 'POST':
        if form.validate_on_submit():
            new_support = ProjectProposalDevelopmentSupport()
            form.populate_obj(new_support)
            new_support.qualification = '|'.join(form.qualification_select.data)
            new_support.docs = '|'.join(form.docs_select.data)
            new_support.project_id = project_id
            new_support.submitted_at = arrow.now(tz='Asia/Bangkok').datetime,
            if form.contract_upload.data:
                upfile = form.contract_upload.data
                filename = secure_filename(upfile.filename)
                upfile.save(filename)
                file_drive = drive.CreateFile({'title': filename})
                file_drive.SetContentFile(filename)
                file_drive.Upload()
                permission = file_drive.InsertPermission({'type': 'anyone',
                                                          'value': 'anyone',
                                                          'role': 'reader'})
                new_support.contract_file_url = file_drive['id']
            db.session.add(new_support)
            db.session.commit()
            flash('New proposal development support request added.')
            return redirect(url_for('project.display_project', project_id=project_id))
    return render_template('project/proposal_development_support_add.html',
                           project_id=project_id, form=form)


@project.route('/<int:project_id>/export')
@login_required
def export_docx(project_id):
    def generate():
        project = ProjectRecord.query.get(project_id)
        doc = Document()
        members = []
        for member in project.members:
            if member.user:
                members.append('{} ({})'.format(member.user.profile.fullname_th, member.role))
            else:
                members.append('{} ({})'.format(member.fullname, member.affil))
        doc.add_heading('รายละเอียดโครงการวิจัย', 0)
        doc.add_heading('ชื่อภาษาไทย', level=1)
        doc.add_heading(project.title_th, level=1)
        doc.add_heading(project.subtitle_th, level=2)
        doc.add_heading('ชื่อภาษาอังกฤษ', level=1)
        doc.add_heading(project.title_en, level=3)
        doc.add_heading(project.subtitle_en, level=4)
        doc.add_heading('แหล่งทุน', level=1)
        doc.add_paragraph(project.fund_source.name)
        doc.add_heading('ที่ปรึกษาโครงการ', level=1)
        doc.add_paragraph(project.mentor)
        doc.add_heading('คณะผู้วิจัย', level=1)
        doc.add_paragraph(', '.join(members))
        doc.add_heading('บทคัดย่อ', level=1)
        doc.add_paragraph(project.abstract)
        doc.add_heading('บทนำ', level=1)
        doc.add_paragraph(project.intro)
        doc.add_heading('วัตถุประสงค์', level=1)
        doc.add_paragraph(project.objective)
        doc.add_heading('ระเบียบวิธีวิจัย', level=1)
        doc.add_paragraph(project.method)
        doc.add_heading('วารสารที่คาดว่าจะเผยแพร่', level=1)
        doc.add_paragraph(project.prospected_journals)
        doc.add_heading('การนำไปใช้ประโยชน์', level=1)
        doc.add_paragraph(project.use_applications)
        doc.add_heading('แก้ไขล่าสุดเมื่อวันที่', level=1)
        doc.add_paragraph(project.updated_at.strftime('วันที่ %d-%m-%Y เวลา %H:%M นาฬิกา'))
        doc.add_heading('ผู้ส่งโครงการ', level=1)
        doc.add_paragraph('{} {}'.format(current_user.profile.title_th, current_user.profile.fullname_th))
        doc.add_heading('สถานะโครงการ', level=1)
        doc.add_paragraph(project.status)
        out_stream = BytesIO()
        doc.save(out_stream)
        return out_stream.getvalue()

    return Response(generate(),
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@project.route('/<int:project_id>/reviews')
@login_required
def view_reviews(project_id):
    project = ProjectRecord.query.get(project_id)
    review = ProjectReviewRecord.query.filter_by(
        project_id=project.id, summarized=True).first()
    return render_template('project/summarized_review.html', project=project, review=review)


@project.route('/<int:project_id>/supplementary', methods=['GET', 'POST'])
@login_required
def add_supplementary_doc(project_id):
    project = ProjectRecord.query.get(project_id)
    form = SupplementaryDocumentForm()
    if request.method == 'POST':
        sup = ProjectSupplementaryDocument()
        form.populate_obj(sup)
        sup.project = project
        if form.file_upload.data:
            upfile = form.file_upload.data
            filename = secure_filename(upfile.filename)
            upfile.save(filename)
            file_drive = drive.CreateFile({'title': filename})
            file_drive.SetContentFile(filename)
            file_drive.Upload()
            permission = file_drive.InsertPermission({'type': 'anyone',
                                                      'value': 'anyone',
                                                      'role': 'reader'})
            sup.file_url = file_drive['id']
            db.session.add(sup)
            db.session.commit()
            flash('ไฟล์ได้รับการบันทึกในระบบแล้ว', 'success')
        return redirect(url_for('project.display_project', project_id=project.id))
    return render_template('project/supplementary_form.html', form=form, project=project)


@project.route('/<int:project_id>/supplementary/<int:doc_id>/remove')
@login_required
def remove_supplementary_doc(project_id, doc_id):
    doc = ProjectSupplementaryDocument.query.get(doc_id)
    file1 = drive.CreateFile({'id': doc.file_url})
    file1.Trash()
    db.session.delete(doc)
    db.session.commit()
    flash('ลบไฟล์ออกจากระบบเรียบร้อยแล้ว', 'success')
    return redirect(url_for('project.display_project', project_id=project_id))



@project.route('/<int:project_id>/overall_budgets/add', methods=['GET', 'POST'])
@login_required
def add_overall_budget_item(project_id):
    project = ProjectRecord.query.get(project_id)
    form = ProjectBudgetItemOverallForm()
    if request.method == 'POST':
        item = ProjectBudgetItemOverall()
        form.populate_obj(item)
        item.project = project
        item.created_at = arrow.now(tz='Asia/Bangkok').datetime
        item.edited_at = arrow.now(tz='Asia/Bangkok').datetime
        db.session.add(item)
        db.session.commit()
        flash('รายการงบการเงินได้รับการบันทึกเรียบร้อยแล้ว', 'success')
        return redirect(url_for('project.display_project', project_id=project.id))
    return render_template('project/overall_budget_form.html', form=form, project=project)


@project.route('/<int:project_id>/overall_budgets/<int:item_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_overall_budget_item(item_id, project_id):
    project = ProjectRecord.query.get(project_id)
    item = ProjectBudgetItemOverall.query.get(item_id)
    form = ProjectBudgetItemOverallForm(obj=item)
    if request.method == 'POST':
        form.populate_obj(item)
        item.project = project
        item.edited_at = arrow.now(tz='Asia/Bangkok').datetime
        db.session.add(item)
        db.session.commit()
        flash('ลบรายการงบการเงินเรียบร้อยแล้ว', 'success')
        return redirect(url_for('project.display_project', project_id=project.id))
    return render_template('project/overall_budget_form.html', form=form, project=project)


@project.route('/<int:project_id>/overall_budgets/<int:item_id>/remove', methods=['GET', 'POST'])
@login_required
def remove_overall_budget_item(item_id, project_id):
    project = ProjectRecord.query.get(project_id)
    item = ProjectBudgetItemOverall.query.get(item_id)
    if item:
        db.session.delete(item)
        db.session.commit()
        flash('รายการงบการเงินได้รับการบันทึกเรียบร้อยแล้ว', 'success')
    else:
        flash('ไม่พบรายการในฐานข้อมูลของระบบ', 'warning')
    return redirect(url_for('project.display_project', project_id=project.id))


@project.route('/<int:project_id>/activity/add', methods=['GET', 'POST'])
@login_required
def add_overall_gantt_activity(project_id):
    project = ProjectRecord.query.get(project_id)
    form = ProjectOverallGanttActivityForm()
    if request.method == 'POST':
        new_activity = ProjectOverallGanttActivity()
        form.populate_obj(new_activity)
        new_activity.project = project
        db.session.add(new_activity)
        db.session.commit()
        flash('เพิ่มกิจกรรมใหม่เรียบร้อย', 'success')
        return redirect(url_for('project.display_project', project_id=project.id))
    return render_template('project/overall_gantt_form.html', form=form, project=project)


@project.route('/activity/<int:activity_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_overall_gantt_activity(activity_id):
    activity = ProjectOverallGanttActivity.query.get(activity_id)
    form = ProjectOverallGanttActivityForm(obj=activity)
    if request.method == 'POST':
        form.populate_obj(activity)
        db.session.add(activity)
        db.session.commit()
        flash('แก้ไขกิจกรรมใหม่เรียบร้อย', 'success')
        return redirect(url_for('project.display_project', project_id=activity.project.id))
    return render_template('project/overall_gantt_form.html', form=form, project=activity.project)


@project.route('/activity/<int:activity_id>/remove', methods=['GET', 'POST'])
@login_required
def remove_overall_gantt_activity(activity_id):
    activity = ProjectOverallGanttActivity.query.get(activity_id)
    project_id = activity.project.id
    if activity:
        db.session.delete(activity)
        db.session.commit()
        flash('ลบกิจกรรมเรียบร้อย', 'success')
    else:
        flash('ไม่พบกิจกรรมในฐานข้อมูล', 'warning')
    return redirect(url_for('project.display_project', project_id=project_id))


@project.route('/<int:project_id>/file_upload', methods=['GET', 'POST'])
@login_required
def upload_summary_file(project_id):
    project = ProjectRecord.query.get(project_id)
    form = FileUploadForm()
    if request.method == 'POST':
        if form.file_upload.data:
            upfile = form.file_upload.data
            filename = secure_filename(upfile.filename)
            upfile.save(filename)
            file_drive = drive.CreateFile({'title': filename})
            file_drive.SetContentFile(filename)
            file_drive.Upload()
            permission = file_drive.InsertPermission({'type': 'anyone',
                                                      'value': 'anyone',
                                                      'role': 'reader'})
            if form.file_type.data == 'finance':
                project.finance_summary_file_url = file_drive['id']
            elif form.file_type.data == 'bookbank_cover':
                project.bookbank_cover_file_url = file_drive['id']
            elif form.file_type.data == 'bookbank_last_page':
                project.bookbank_last_page_file_url = file_drive['id']
            db.session.add(project)
            db.session.commit()
            flash('อัพโหลดไฟล์เรียบร้อย', 'success')
            return redirect(url_for('project.display_project', project_id=project.id))
    return render_template('project/file_upload.html', form=form)
